import os
import random
import re

from win32com import client as wc
from docxcompose.image import ImageWrapper
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docxcompose.composer import Composer
from docxcompose.utils import xpath
from docxcompose.utils import NS
from copy import deepcopy
from docx import Document


def add_shapes(self, doc, element):
    shapes = xpath(element, './/v:shape/v:imagedata')
    for shape in shapes:
        rid = shape.get('{%s}id' % NS['r'])
        if rid is None:  # 这里会有时候取不到 导致报错 重写
            continue
        img_part = doc.part.rels[rid].target_part

        new_img_part = self.pkg.image_parts._get_by_sha1(img_part.sha1)
        if new_img_part is None:
            image = ImageWrapper(img_part)
            new_img_part = self.pkg.image_parts._add_image_part(image)

        new_rid = self.doc.part.relate_to(new_img_part, RT.IMAGE)
        shape.set('{%s}id' % NS['r'], new_rid)


Composer.add_shapes = add_shapes


def add_numberings(self, doc, element):
    """Add numberings from the given document used in the given element."""
    # Search for numbering references
    num_ids = set([n.val for n in xpath(element, './/w:numId')])
    if not num_ids:
        return

    next_num_id, next_anum_id = self._next_numbering_ids()
    try:
        src_numbering_part = doc.part.numbering_part
    except NotImplementedError:
        src_numbering_part = self.numbering_part()

    for num_id in num_ids:
        if num_id in self.num_id_mapping:
            continue

        # Find the referenced <w:num> element
        res = src_numbering_part.element.xpath(
            './/w:num[@w:numId="%s"]' % num_id)
        if not res:
            continue
        num_element = deepcopy(res[0])
        num_element.numId = next_num_id

        self.num_id_mapping[num_id] = next_num_id

        anum_id = num_element.xpath('//w:abstractNumId')[0]
        if anum_id.val not in self.anum_id_mapping:
            # Find the referenced <w:abstractNum> element
            res = src_numbering_part.element.xpath(
                './/w:abstractNum[@w:abstractNumId="%s"]' % anum_id.val)
            if not res:
                continue
            anum_element = deepcopy(res[0])
            self.anum_id_mapping[anum_id.val] = next_anum_id
            anum_id.val = next_anum_id
            # anum_element.abstractNumId = next_anum_id
            anum_element.set('{%s}abstractNumId' % NS['w'], str(next_anum_id))

            # Make sure we have a unique nsid so numberings restart properly
            nsid = anum_element.find('.//w:nsid', NS)
            if nsid is not None:
                nsid.set(
                    '{%s}val' % NS['w'],
                    "{0:08X}".format(int(10 ** 8 * random.random()))
                )

            self._insert_abstract_num(anum_element)
        else:
            anum_id.val = self.anum_id_mapping[anum_id.val]

        self._insert_num(num_element)

    # Fix references
    for num_id_ref in xpath(element, './/w:numId'):
        num_id_ref.val = self.num_id_mapping.get(
            num_id_ref.val, num_id_ref.val)


Composer.add_numberings = add_numberings


def merge_word_documents(folder_path, output_path):
    # ---------- 读取文件目录 -------------
    cwd = os.getcwd()
    file_dir_list = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.endswith('.doc'):
                file_dir_list.append(os.path.join(cwd, root, file))
    print(f'一共找到{len(file_dir_list)}个doc文件')

    # ------ 将doc另存为docx ----------
    word = wc.Dispatch("Word.Application")
    for file in file_dir_list:
        doc = word.Documents.Open(file)
        doc.SaveAs(f'{file}x', 12)
        doc.Close()
    word.Quit()

    # ---------- 读取docx文件路径 ----------
    docx_list = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.endswith('.docx'):
                docx_list.append(os.path.join(cwd, root, file))
    print(f'一共找到{len(docx_list)}个docx文件')

    # -------- 按照顺序排列，这步不同人的电脑上可能会有报错，需要自己调试 -------------
    # -------- 此处正则'^[-+]?\d+(\.\d+)?'为匹配文件名开头的整数或小数 -------------
    # -------- 如: 3 观潮.docx 中的 3 -------------
    # -------- 如: 5.2 爬山虎的脚.docx 中的 5.2 -------------
    # -------- l[l.rfind(os.sep) + 1:] 为截取取文件名-------------
    # -------- (以windows举例)l原始为: D:\python\mergeWordOrigin\allWord\1 观潮教案.docx -------------
    # -------- os.sep 为当前系统路径的分隔符 -------------
    # -------- l[l.rfind(os.sep) + 1:] 为从l中截取最后一个斜杠之后的内容 -------------
    # -------- .group(0) 为取正则匹配到的第一个字符串 -------------
    docx_list.sort(key=lambda l: float(re.search(r'^[-+]?\d+(\.\d+)?', l[l.rfind(os.sep) + 1:]).group(0)))
    for each in docx_list:
        print(each)

    # -------- 合并文档 ---------
    # 填充分页符号文档
    page_break_doc = Document()
    page_break_doc.add_page_break()
    # 定义新文档
    target_doc = Document(docx_list[0])
    target_composer = Composer(target_doc)
    for i in range(len(docx_list)):
        # 跳过第一个作为模板的文件
        if i == 0:
            continue
        # 填充分页符文档
        target_composer.append(page_break_doc)
        # 拼接文档内容
        f = docx_list[i]
        target_composer.append(Document(f))
    # 保存目标文档
    target_composer.save(output_path)


# 程序入口
if __name__ == "__main__":
    folder_path = "allWord"
    output_path = "all.docx"
    merge_word_documents(folder_path, output_path)
